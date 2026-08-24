from __future__ import annotations

import hashlib
from pathlib import Path

from docx import Document

from .contracts import EvidenceAtom, FactObservation
from .text_extraction import extract_facts_from_text
from .xlsx_parser import load_rules, sha256_file


def parse_docx(path: Path, root: Path, rules_path: Path) -> tuple[list[EvidenceAtom], list[FactObservation]]:
    rules = load_rules(rules_path)
    document_hash = sha256_file(path)
    relative_path = path.relative_to(root).as_posix()
    document = Document(path)
    evidence: list[EvidenceAtom] = []
    observations: list[FactObservation] = []

    for index, paragraph in enumerate(document.paragraphs, start=1):
        text = paragraph.text.strip()
        if text:
            _append_block(text, f"paragraph:{index}", document_hash, relative_path, rules, evidence, observations)

    for table_index, table in enumerate(document.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            for column_index, cell in enumerate(row.cells, start=1):
                text = cell.text.strip()
                if not text:
                    continue
                label = row.cells[column_index - 2].text.strip() if column_index > 1 else ""
                block = f"{label}: {text}" if label else text
                locator = f"table:{table_index}/row:{row_index}/col:{column_index}"
                _append_block(block, locator, document_hash, relative_path, rules, evidence, observations, raw_text=text)

    return evidence, observations


def _append_block(
    text: str,
    locator: str,
    document_hash: str,
    relative_path: str,
    rules: dict,
    evidence: list[EvidenceAtom],
    observations: list[FactObservation],
    raw_text: str | None = None,
) -> None:
    evidence_id = hashlib.sha256(f"{document_hash}|{locator}".encode("utf-8")).hexdigest()[:20]
    evidence.append(
        EvidenceAtom(
            evidence_id=evidence_id,
            document_hash=document_hash,
            relative_path=relative_path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            locator=locator,
            raw_text=raw_text or text,
            extractor="python-docx",
        )
    )
    observations.extend(
        extract_facts_from_text(
            text=text,
            base_locator=locator,
            document_hash=document_hash,
            relative_path=relative_path,
            evidence_id=evidence_id,
            rules=rules,
        )
    )

